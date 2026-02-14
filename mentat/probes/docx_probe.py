import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from mentat.probes.base import (
    BaseProbe,
    ProbeResult,
    TopicInfo,
    StructureInfo,
    TocEntry,
    Chunk,
)
from mentat.probes._utils import estimate_tokens, should_bypass, extract_preview, merge_small_chunks
from mentat.probes.instruction_templates import (
    DOCX_BRIEF_INTRO,
    DOCX_INSTRUCTIONS,
)

try:
    from docx import Document

    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

_HEADING_STYLE_RE = re.compile(r"Heading\s+(\d+)", re.IGNORECASE)
_LIST_STYLE_RE = re.compile(r"List", re.IGNORECASE)


class DOCXProbe(BaseProbe):
    """Probe for Word documents (.docx)."""

    def can_handle(self, filename: str, content_type: str) -> bool:
        if not _DOCX_AVAILABLE:
            return False
        return filename.lower().endswith(".docx")

    def run(self, file_path: str) -> ProbeResult:
        doc = Document(file_path)

        # --- Extract all text for token estimation ---
        all_text_parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_parts.append(para.text)
        full_text = "\n".join(all_text_parts)
        approx_tokens = estimate_tokens(full_text)

        # --- Document metadata ---
        props = doc.core_properties
        meta: Dict[str, Any] = {}
        if props.title:
            meta["title"] = props.title
        if props.author:
            meta["author"] = props.author
        if props.subject:
            meta["subject"] = props.subject
        if props.created:
            meta["created"] = str(props.created)
        if props.modified:
            meta["modified"] = str(props.modified)
        if props.revision:
            meta["revision"] = props.revision

        # --- Build heading hierarchy with preview/annotation ---
        toc_entries, sections = self._extract_headings(doc)

        # --- Table detection ---
        table_count = len(doc.tables)
        table_headers: List[str] = []
        for table in doc.tables[:5]:
            if table.rows:
                header_row = [cell.text.strip() for cell in table.rows[0].cells]
                table_headers.append(" | ".join(header_row))

        # --- Stats ---
        stats: Dict[str, Any] = {
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(full_text.split()),
            "approx_tokens": approx_tokens,
            "table_count": table_count,
            **meta,
        }
        if table_headers:
            stats["table_headers"] = table_headers

        # --- Topic ---
        doc_title = meta.get("title") or (
            toc_entries[0].title if toc_entries else Path(file_path).stem
        )
        first_para = None
        for para in doc.paragraphs:
            if para.text.strip() and not _HEADING_STYLE_RE.match(para.style.name or ""):
                first_para = para.text.strip()[:300]
                break

        topic = TopicInfo(
            title=doc_title,
            abstract=meta.get("subject"),
            first_paragraph=first_para,
        )

        structure = StructureInfo(toc=toc_entries)

        # --- Small-file bypass ---
        if should_bypass(full_text):
            stats["is_full_content"] = True
            result = ProbeResult(
                filename=Path(file_path).name,
                file_type="docx",
                topic=topic,
                structure=structure,
                stats=stats,
                chunks=[Chunk(content=full_text, index=0)],
                raw_snippet=full_text,
            )
            brief_intro, instructions = self.generate_instructions(result)
            result.brief_intro = brief_intro
            result.instructions = instructions
            return result

        # --- Chunks: split by headings ---
        stats["is_full_content"] = False
        chunks = self._build_chunks(sections)

        result = ProbeResult(
            filename=Path(file_path).name,
            file_type="docx",
            topic=topic,
            structure=structure,
            stats=stats,
            chunks=chunks,
            raw_snippet=full_text[:500],
        )

        # Generate format-specific instructions
        brief_intro, instructions = self.generate_instructions(result)
        result.brief_intro = brief_intro
        result.instructions = instructions

        return result

    def generate_instructions(self, probe_result: ProbeResult) -> Tuple[str, str]:
        """Generate DOCX-specific instructions."""
        # Brief intro
        brief_intro = DOCX_BRIEF_INTRO

        # Full instructions
        instructions = DOCX_INSTRUCTIONS.format(filename=probe_result.filename)

        return brief_intro, instructions

    def _extract_headings(
        self, doc: "Document"
    ) -> tuple:
        """Extract heading hierarchy with preview and annotation."""
        toc_entries: List[TocEntry] = []
        sections: List[Dict[str, Any]] = []  # {title, level, paragraphs}

        current_section: Optional[Dict[str, Any]] = None

        for para in doc.paragraphs:
            style_name = para.style.name or ""
            heading_match = _HEADING_STYLE_RE.match(style_name)

            if heading_match:
                # Save previous section
                if current_section:
                    sections.append(current_section)
                    self._finalize_section(current_section, toc_entries)

                level = int(heading_match.group(1))
                current_section = {
                    "title": para.text.strip(),
                    "level": level,
                    "paragraphs": [],
                    "list_items": 0,
                }
            elif current_section is not None:
                text = para.text.strip()
                if text:
                    current_section["paragraphs"].append(text)
                    if _LIST_STYLE_RE.search(style_name):
                        current_section["list_items"] += 1
            else:
                # Content before first heading
                if para.text.strip() and not sections:
                    if current_section is None:
                        current_section = {
                            "title": "Preamble",
                            "level": 0,
                            "paragraphs": [],
                            "list_items": 0,
                        }
                    current_section["paragraphs"].append(para.text.strip())

        # Finalize last section
        if current_section:
            sections.append(current_section)
            self._finalize_section(current_section, toc_entries)

        return toc_entries, sections

    def _finalize_section(
        self, section: Dict[str, Any], toc_entries: List[TocEntry]
    ):
        """Build a TocEntry from a completed section."""
        parts = []
        para_count = len(section["paragraphs"])
        if section["list_items"]:
            parts.append(f"List, {section['list_items']} items")
        parts.append(f"{para_count} paragraphs")
        annotation = " | ".join(parts)

        preview = section["paragraphs"][0][:120] if section["paragraphs"] else None

        toc_entries.append(
            TocEntry(
                level=max(section["level"], 1),
                title=section["title"],
                annotation=annotation,
                preview=preview,
            )
        )

    def _build_chunks(self, sections: List[Dict[str, Any]]) -> List[Chunk]:
        chunks: List[Chunk] = []
        for section in sections:
            text = section["title"] + "\n\n" + "\n".join(section["paragraphs"])
            level = section.get("level", 1)
            chunks.append(
                Chunk(
                    content=text.strip(),
                    index=len(chunks),
                    section=section["title"],
                    metadata={"level": level},
                )
            )
        if not chunks:
            return [Chunk(content="(empty document)", index=0)]
        return merge_small_chunks(chunks)
